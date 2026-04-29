"""
document_processor.py
─────────────────────
Handles text and image extraction from uploaded PDF / DOCX / DOC files
for the CTI-to-STIX pipeline.

Exports
-------
process_document(file_bytes, filename, image_output_dir)
    -> (cleaned_text: str, image_paths: list[str], status: str)
"""

import os
import re
import io
import unicodedata


# ──────────────────────────────────────────────────────────────────────────────
# TEXT CLEANING
# ──────────────────────────────────────────────────────────────────────────────

def clean_text(text: str) -> str:
    """
    Normalise and clean raw extracted text so it is suitable for the
    downstream CTI/STIX entity-extraction pipeline.

    Preserves cybersecurity entities (IPs, domains, hashes, CVEs, ATT&CK
    IDs, e-mail addresses, file names, URLs) while stripping formatting
    artefacts introduced by PDF/Word renderers.
    """
    # 1. Unicode NFC normalisation
    text = unicodedata.normalize("NFC", text)

    # 2. Smart / curly quotes → ASCII
    text = (text
            .replace("\u2018", "'").replace("\u2019", "'")   # ' '
            .replace("\u201c", '"').replace("\u201d", '"')   # " "
            .replace("\u201a", ",").replace("\u201e", '"'))  # ‚ „

    # 3. Dashes → ASCII hyphen (keeps CVE IDs and IP ranges intact)
    text = (text
            .replace("\u2014", " - ")   # em dash
            .replace("\u2013", " - ")   # en dash
            .replace("\u2012", "-")     # figure dash
            .replace("\u2015", "-"))    # horizontal bar

    # 4. Common typographic ligatures → ASCII sequences
    for lig, rep in {
        "\ufb00": "ff", "\ufb01": "fi", "\ufb02": "fl",
        "\ufb03": "ffi", "\ufb04": "ffl", "\ufb06": "st",
    }.items():
        text = text.replace(lig, rep)

    # 5. Invisible / special whitespace variants
    text = (text
            .replace("\u00a0", " ")   # non-breaking space
            .replace("\u200b", "")    # zero-width space
            .replace("\u200c", "")    # zero-width non-joiner
            .replace("\u200d", "")    # zero-width joiner
            .replace("\ufeff", ""))   # BOM

    # 6. Decorative bullet characters → "- " (keeps list items readable)
    for bullet in "•◦▪▸►▶→✓✗✘★☆◆◇▷◁":
        text = text.replace(bullet, "- ")

    # 7. Line-ending normalisation + page-break removal
    text = text.replace("\r\n", "\n").replace("\r", "\n").replace("\f", "\n")

    # 8. Strip stray ASCII control characters (keep \t and \n)
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # 9. Rejoin words split by a hyphen + line-break (common in PDF columns)
    #    e.g. "infra-\nstructure" → "infrastructure"
    #    Only rejoins when both sides are word characters (won't touch CVE-2021-\n40444
    #    because the right side starts with a digit, not a letter — safe enough).
    text = re.sub(r"([a-zA-Z])-\n([a-zA-Z])", r"\1\2", text)

    # 10. Collapse horizontal whitespace within lines
    text = re.sub(r"[ \t]+", " ", text)

    # 11. Per-line strip + remove purely decorative separator lines
    lines = []
    for line in text.split("\n"):
        stripped = line.strip()
        # Drop lines that are only dashes, equals, underscores etc.
        if re.fullmatch(r"[-=_*#~+]{5,}", stripped):
            continue
        lines.append(stripped)

    # 12. Collapse 3+ consecutive blank lines → 2 (preserves paragraph breaks)
    cleaned_lines: list[str] = []
    blank_run = 0
    for line in lines:
        if line == "":
            blank_run += 1
            if blank_run <= 2:
                cleaned_lines.append(line)
        else:
            blank_run = 0
            cleaned_lines.append(line)

    text = "\n".join(cleaned_lines).strip()

    # 13. Re-fang defanged IOCs that CTI authors write defensively
    #     e.g. evil[.]com → evil.com  |  hxxp:// → http://  |  192[.]168 → 192.168
    #     Uses ioc_fanger if available; silently skips if not installed.
    try:
        from ioc_fanger import ioc_fanger as fanger
        text = fanger.fang(text)
    except Exception:
        pass

    return text


# ──────────────────────────────────────────────────────────────────────────────
# PDF EXTRACTION  (PyMuPDF)
# ──────────────────────────────────────────────────────────────────────────────

def extract_from_pdf(file_bytes: bytes, image_output_dir: str):
    """
    Extract text and images from a PDF using PyMuPDF (fitz).
    Returns (cleaned_text, image_paths, status).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return "", [], "PyMuPDF is not installed. Run: pip install PyMuPDF"

    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:
        return "", [], f"Could not open PDF: {exc}"

    if doc.is_encrypted:
        return "", [], (
            "PDF is password-protected / encrypted. "
            "Please decrypt it before uploading."
        )
    if doc.page_count == 0:
        return "", [], "PDF has no pages."

    text_parts: list[str] = []
    image_paths: list[str] = []
    seen_xrefs: set[int] = set()  # deduplicate shared images across pages

    for page_num in range(doc.page_count):
        page = doc[page_num]

        # ── text ──
        page_text = page.get_text("text")
        if page_text.strip():
            text_parts.append(page_text)

        # ── images ──
        for img_index, img_info in enumerate(page.get_images(full=True)):
            xref = img_info[0]
            if xref in seen_xrefs:
                continue
            seen_xrefs.add(xref)
            try:
                base_image = doc.extract_image(xref)
                img_bytes = base_image["image"]
                img_ext   = base_image.get("ext", "png")
                # Skip tiny images (icons, bullets, decorative glyphs)
                if len(img_bytes) < 4096:
                    continue
                img_name = f"pdf_p{page_num + 1}_img{img_index + 1}.{img_ext}"
                img_path = os.path.join(image_output_dir, img_name)
                with open(img_path, "wb") as fh:
                    fh.write(img_bytes)
                image_paths.append(img_path)
            except Exception:
                pass  # skip individual image failures

    doc.close()

    raw_text = "\n\n".join(text_parts)
    cleaned  = clean_text(raw_text)

    if not cleaned.strip():
        status = (
            "⚠️ Minimal text extracted — the PDF may be scanned / image-only. "
            "OCR is not supported; please use a text-based PDF."
        )
    else:
        status = "OK"

    return cleaned, image_paths, status


# ──────────────────────────────────────────────────────────────────────────────
# DOCX EXTRACTION  (python-docx)
# ──────────────────────────────────────────────────────────────────────────────

def extract_from_docx(file_bytes: bytes, image_output_dir: str):
    """
    Extract text and images from a DOCX file using python-docx.
    Tables are rendered as pipe-separated rows so they remain readable.
    Returns (cleaned_text, image_paths, status).
    """
    try:
        from docx import Document
    except ImportError:
        return "", [], "python-docx is not installed. Run: pip install python-docx"

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        return "", [], f"Could not open DOCX: {exc}"

    text_parts: list[str] = []

    # ── paragraphs ──
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # ── tables (flatten to readable plain text) ──
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Deduplicate merged cells (python-docx repeats them)
            deduped: list[str] = []
            for cell in cells:
                if not deduped or cell != deduped[-1]:
                    deduped.append(cell)
            row_str = " | ".join(deduped)
            if row_str.strip():
                text_parts.append(row_str)

    # ── images ──
    image_paths: list[str] = []
    try:
        for rel in doc.part.rels.values():
            if "image" not in rel.reltype.lower():
                continue
            try:
                img_data = rel.target_part.blob
                ct  = rel.target_part.content_type          # e.g. "image/png"
                ext = ct.split("/")[-1].replace("jpeg", "jpg")
                idx = len(image_paths) + 1
                img_name = f"docx_img{idx}.{ext}"
                img_path = os.path.join(image_output_dir, img_name)
                with open(img_path, "wb") as fh:
                    fh.write(img_data)
                image_paths.append(img_path)
            except Exception:
                pass
    except Exception:
        pass

    raw_text = "\n".join(text_parts)
    cleaned  = clean_text(raw_text)
    status   = "OK" if cleaned.strip() else "⚠️ No text could be extracted from the DOCX."
    return cleaned, image_paths, status


# ──────────────────────────────────────────────────────────────────────────────
# LEGACY .DOC EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────

def extract_from_doc(file_bytes: bytes):
    """
    Best-effort text extraction from a legacy binary .doc file.
    Image extraction is not supported for this format.
    Returns (cleaned_text, [], status).
    """
    import tempfile

    # Try docx2txt — handles some .doc files without external tools
    try:
        import docx2txt
        with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            text = docx2txt.process(tmp_path)
            if text and text.strip():
                return clean_text(text), [], "OK (note: .doc image extraction not supported)"
        except Exception:
            pass
        finally:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    except ImportError:
        pass

    return "", [], (
        "⚠️ Could not extract text from this legacy .doc file. "
        "Please re-save it as .docx in Microsoft Word or LibreOffice and re-upload."
    )


# ──────────────────────────────────────────────────────────────────────────────
# PUBLIC ENTRY POINT
# ──────────────────────────────────────────────────────────────────────────────

def process_document(file_bytes: bytes, filename: str, image_output_dir: str):
    """
    Dispatch to the appropriate extractor based on file extension.

    Parameters
    ----------
    file_bytes       : raw bytes of the uploaded file
    filename         : original filename (used for extension detection)
    image_output_dir : directory where extracted images will be saved

    Returns
    -------
    (extracted_text: str, image_paths: list[str], status: str)
        status == "OK" on success; a warning/error string otherwise.
    """
    os.makedirs(image_output_dir, exist_ok=True)
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return extract_from_pdf(file_bytes, image_output_dir)
    elif ext == ".docx":
        return extract_from_docx(file_bytes, image_output_dir)
    elif ext == ".doc":
        return extract_from_doc(file_bytes)
    else:
        return "", [], (
            f"Unsupported file type '{ext}'. "
            "Please upload a PDF (.pdf) or Word document (.doc / .docx)."
        )
